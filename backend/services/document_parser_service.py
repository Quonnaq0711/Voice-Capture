"""Unified document parser service for parsing multiple file formats (PDF, DOCX, TXT)."""

import os
import logging
from typing import Optional
from pathlib import Path

logger = logging.getLogger(__name__)


class DocumentParserService:
    """Service for parsing various document formats into plain text."""

    SUPPORTED_FORMATS = ['.pdf', '.docx', '.txt']

    @staticmethod
    def is_supported_format(file_path: str) -> bool:
        """
        Check if the file format is supported.

        Args:
            file_path: Path to the file

        Returns:
            True if format is supported, False otherwise
        """
        ext = Path(file_path).suffix.lower()
        return ext in DocumentParserService.SUPPORTED_FORMATS

    @staticmethod
    def parse_document(file_path: str) -> Optional[str]:
        """
        Parse a document file and extract text content.
        Supports PDF, DOCX, and TXT formats.

        Args:
            file_path: Path to the document file

        Returns:
            Extracted text content as string, or None if parsing fails

        Raises:
            FileNotFoundError: If the file doesn't exist
            ValueError: If the file format is not supported
        """
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            raise FileNotFoundError(f"File not found: {file_path}")

        file_extension = Path(file_path).suffix.lower()

        if file_extension not in DocumentParserService.SUPPORTED_FORMATS:
            logger.error(f"Unsupported file format: {file_extension}")
            raise ValueError(f"Unsupported file format: {file_extension}. Supported formats: {', '.join(DocumentParserService.SUPPORTED_FORMATS)}")

        try:
            if file_extension == '.pdf':
                return DocumentParserService._parse_pdf(file_path)
            elif file_extension == '.docx':
                return DocumentParserService._parse_docx(file_path)
            elif file_extension == '.txt':
                return DocumentParserService._parse_txt(file_path)
        except Exception as e:
            logger.error(f"Error parsing document {file_path}: {str(e)}", exc_info=True)
            return None

    @staticmethod
    def _parse_pdf(file_path: str) -> Optional[str]:
        """
        Parse PDF file and extract text.

        Args:
            file_path: Path to the PDF file

        Returns:
            Extracted text content
        """
        try:
            import PyPDF2

            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                content = ""

                for page_num, page in enumerate(reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            content += page_text + "\n"
                    except Exception as e:
                        logger.warning(f"Error extracting text from page {page_num + 1} in {file_path}: {str(e)}")
                        continue

                if not content.strip():
                    logger.warning(f"No text content extracted from PDF: {file_path}")
                    return None

                logger.info(f"Successfully parsed PDF: {file_path} ({len(reader.pages)} pages)")
                return content.strip()

        except Exception as e:
            logger.error(f"Error parsing PDF file {file_path}: {str(e)}")
            return None

    @staticmethod
    def _parse_docx(file_path: str) -> Optional[str]:
        """
        Parse Word document (DOCX) and extract text.
        Handles both standard and non-standard DOCX formats.

        Args:
            file_path: Path to the DOCX document

        Returns:
            Extracted text content
        """
        try:
            from docx import Document

            doc = Document(file_path)
            content = []

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content.append(paragraph.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            content.append(cell.text)

            if not content:
                logger.warning(f"No text content extracted from Word document: {file_path}")
                return None

            full_text = "\n".join(content)
            logger.info(f"Successfully parsed Word document: {file_path} ({len(doc.paragraphs)} paragraphs)")
            return full_text.strip()

        except KeyError as e:
            # Handle non-standard DOCX formats (e.g., from online converters)
            if "officeDocument" in str(e):
                logger.warning(f"Non-standard DOCX format detected in {file_path}, trying alternative parsing method")
                return DocumentParserService._parse_docx_alternative(file_path)
            else:
                logger.error(f"Error parsing Word document {file_path}: {str(e)}")
                return None
        except Exception as e:
            logger.error(f"Error parsing Word document {file_path}: {str(e)}")
            # Try alternative parsing method as fallback
            try:
                return DocumentParserService._parse_docx_alternative(file_path)
            except:
                return None

    @staticmethod
    def _parse_docx_alternative(file_path: str) -> Optional[str]:
        """
        Alternative method to parse DOCX using direct XML extraction.
        Used for non-standard DOCX formats that python-docx cannot handle.

        Args:
            file_path: Path to the DOCX document

        Returns:
            Extracted text content
        """
        try:
            import zipfile
            import xml.etree.ElementTree as ET

            # DOCX is a ZIP file containing XML
            with zipfile.ZipFile(file_path, 'r') as zip_ref:
                # Read the main document XML
                try:
                    doc_xml = zip_ref.read('word/document.xml')
                except KeyError:
                    logger.error(f"No word/document.xml found in {file_path}")
                    return None

                # Parse XML and extract text
                root = ET.fromstring(doc_xml)

                # Try multiple namespace approaches
                text_elements = []

                # Approach 1: Try standard Microsoft namespace
                namespaces_standard = {
                    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
                }
                text_elements = root.findall('.//w:t', namespaces_standard)

                # Approach 2: Try alternative OOXML namespace
                if not text_elements:
                    namespaces_alt = {
                        'w': 'http://purl.oclc.org/ooxml/wordprocessingml/main',
                    }
                    text_elements = root.findall('.//w:t', namespaces_alt)

                # Approach 3: Iterate through all elements and find 't' tags (namespace-agnostic)
                if not text_elements:
                    text_elements = []
                    for elem in root.iter():
                        # Check if tag ends with '}t' (namespace-agnostic way)
                        if elem.tag.endswith('}t') or elem.tag == 't':
                            text_elements.append(elem)

                # Extract text content
                content = []
                for elem in text_elements:
                    if elem.text and elem.text.strip():
                        content.append(elem.text)

                if not content:
                    logger.warning(f"No text content extracted from Word document using alternative method: {file_path}")
                    return None

                full_text = " ".join(content)
                logger.info(f"Successfully parsed Word document using alternative method: {file_path} ({len(text_elements)} text elements)")
                return full_text.strip()

        except Exception as e:
            logger.error(f"Error in alternative DOCX parsing for {file_path}: {str(e)}")
            return None

    @staticmethod
    def _parse_txt(file_path: str) -> Optional[str]:
        """
        Parse plain text file.

        Args:
            file_path: Path to the text file

        Returns:
            File content as string
        """
        try:
            # Try multiple encodings
            encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']

            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()

                    if not content.strip():
                        logger.warning(f"Empty text file: {file_path}")
                        return None

                    logger.info(f"Successfully parsed text file: {file_path} (encoding: {encoding})")
                    return content.strip()

                except UnicodeDecodeError:
                    if encoding == encodings[-1]:
                        # Last encoding failed
                        raise
                    continue

            return None

        except Exception as e:
            logger.error(f"Error parsing text file {file_path}: {str(e)}")
            return None

    @staticmethod
    def get_file_info(file_path: str) -> dict:
        """
        Get basic information about the document.

        Args:
            file_path: Path to the document

        Returns:
            Dictionary containing file information
        """
        try:
            path = Path(file_path)
            stat = path.stat()

            return {
                'filename': path.name,
                'extension': path.suffix.lower(),
                'size_bytes': stat.st_size,
                'size_kb': round(stat.st_size / 1024, 2),
                'is_supported': path.suffix.lower() in DocumentParserService.SUPPORTED_FORMATS
            }
        except Exception as e:
            logger.error(f"Error getting file info for {file_path}: {str(e)}")
            return {}
